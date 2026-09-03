from __future__ import annotations

import datetime as dt
import hashlib
import math
import re
from collections import Counter
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader
from sqlalchemy.orm import Session

from app.core.config import Settings, get_settings
from app.db.models import CurriculumChunkModel, CurriculumSourceModel
from app.services.curriculum_vector_service import (
    CurriculumVectorStore,
    CurriculumVectorUnavailable,
)


SUPPORTED_CURRICULUM_EXTENSIONS = {".md", ".txt", ".pdf", ".docx"}
TOKEN_STOPWORDS = {
    "的",
    "了",
    "和",
    "与",
    "及",
    "在",
    "是",
    "为",
    "对",
    "将",
    "中",
    "可",
    "要",
    "把",
}
EDUCATION_LEVEL_KEYWORDS = {
    "primary": {
        "小学",
        "一年级",
        "二年级",
        "三年级",
        "四年级",
        "五年级",
        "六年级",
    },
    "middle": {"初中", "七年级", "八年级", "九年级"},
    "high": {"高中", "高一", "高二", "高三"},
}
SUBJECT_KEYWORDS = {
    "physics": {
        "物理",
        "力的作用",
        "力学",
        "运动和力",
        "推墙",
        "反作用",
        "机械能",
        "电磁",
    },
    "mathematics": {"数学", "代数", "几何", "函数", "方程", "图形与几何"},
    "information_technology": {
        "信息科技",
        "信息技术",
        "人工智能",
        "编程",
        "python",
        "算法",
    },
}


class CurriculumFileError(ValueError):
    pass


@dataclass(frozen=True)
class CurriculumSearchResult:
    chunk_id: int
    source: str
    source_index: int
    content: str
    score: float
    chunk_ids: tuple[int, ...]
    bm25_score: float = 0.0
    vector_score: float = 0.0
    fusion_score: float = 0.0
    retrieval_mode: str = "local_bm25"


class CurriculumKnowledgeService:
    def __init__(
        self,
        db: Session,
        settings: Settings | None = None,
        vector_store: CurriculumVectorStore | None = None,
    ):
        self.db = db
        self.settings = settings or get_settings()
        self.vector_store = vector_store or CurriculumVectorStore(self.settings)
        self.last_retrieval = {
            "mode": "local_bm25",
            "vector_error": self.vector_store.error,
        }

    def ingest_file(self, path: Path, *, source: str | None = None) -> int:
        path = path.resolve()
        extension = path.suffix.lower()
        if extension not in SUPPORTED_CURRICULUM_EXTENSIONS:
            supported = ", ".join(sorted(SUPPORTED_CURRICULUM_EXTENSIONS))
            raise CurriculumFileError(f"仅支持以下课标文件格式：{supported}")
        if not path.is_file():
            raise CurriculumFileError(f"课标文件不存在：{path}")

        content = self.extract_text(path)
        return self.ingest(source or path.name, content)

    def ingest_bytes(self, filename: str, data: bytes) -> int:
        source = self.safe_source_name(filename)
        content = self.extract_bytes(source, data)
        return self.ingest(source, content)

    def ingest(self, source: str, content: str) -> int:
        normalized_source = source.replace("\\", "/").strip()
        if not normalized_source:
            raise CurriculumFileError("课标来源名称不能为空")

        normalized_content = normalize_text(content)
        if not normalized_content:
            raise CurriculumFileError("课标文件中未提取到可用文字；扫描版 PDF 暂不支持 OCR")

        overlap = min(
            self.settings.curriculum_chunk_overlap,
            self.settings.curriculum_chunk_size - 1,
        )
        chunks = chunk_text(
            normalized_content,
            size=self.settings.curriculum_chunk_size,
            overlap=overlap,
        )
        if not chunks:
            raise CurriculumFileError("课标文件无法切分出有效内容")

        return self.ingest_chunks(normalized_source, chunks)

    def ingest_chunks(
        self,
        source: str,
        chunks: list[str],
        *,
        checksum: str | None = None,
    ) -> int:
        normalized_source = normalize_source(source)
        normalized_chunks = [normalize_text(chunk) for chunk in chunks]
        normalized_chunks = [chunk for chunk in normalized_chunks if chunk]
        if not normalized_chunks:
            raise CurriculumFileError("课标文件无法切分出有效内容")

        try:
            self.vector_store.delete_source(normalized_source)
        except Exception:
            # 数据库是事实源；残留向量在召回时还会经过 chunk_id 校验。
            pass
        self.db.query(CurriculumChunkModel).filter(
            CurriculumChunkModel.source == normalized_source
        ).delete(synchronize_session=False)
        timestamp = now_iso()
        rows = [
            CurriculumChunkModel(
                source=normalized_source,
                source_index=index,
                content=chunk,
                created_at=timestamp,
            )
            for index, chunk in enumerate(normalized_chunks)
        ]
        self.db.add_all(rows)
        self.db.flush()

        vector_count = 0
        vector_status = "disabled"
        vector_error = ""
        if self.settings.curriculum_vector_enabled:
            vector_status = "pending"
            try:
                vector_count = self.vector_store.index_chunks(rows)
                vector_status = "ready"
            except Exception as exc:
                vector_status = "error"
                vector_error = str(exc)
                if self.settings.curriculum_vector_required:
                    raise CurriculumFileError(vector_error) from exc

        source_state = self.db.get(CurriculumSourceModel, normalized_source)
        if source_state is None:
            source_state = CurriculumSourceModel(
                source=normalized_source,
                updated_at=timestamp,
            )
            self.db.add(source_state)
        source_state.checksum = checksum or checksum_chunks(normalized_chunks)
        source_state.chunk_count = len(rows)
        source_state.vector_chunk_count = vector_count
        source_state.vector_status = vector_status
        source_state.embedding_model = (
            self.settings.curriculum_embedding_model
            if self.settings.curriculum_vector_enabled
            else ""
        )
        source_state.last_error = vector_error
        source_state.updated_at = timestamp
        self.db.flush()
        return len(rows)

    def ensure_source_records(self) -> None:
        chunks = (
            self.db.query(CurriculumChunkModel)
            .order_by(CurriculumChunkModel.source, CurriculumChunkModel.source_index)
            .all()
        )
        grouped: dict[str, list[CurriculumChunkModel]] = {}
        for chunk in chunks:
            grouped.setdefault(chunk.source, []).append(chunk)
        timestamp = now_iso()
        for source, rows in grouped.items():
            state = self.db.get(CurriculumSourceModel, source)
            if state is None:
                state = CurriculumSourceModel(source=source, updated_at=timestamp)
                self.db.add(state)
                state.vector_status = (
                    "pending" if self.settings.curriculum_vector_enabled else "disabled"
                )
            state.chunk_count = len(rows)
            if not state.checksum:
                state.checksum = checksum_chunks([row.content for row in rows])
            if not state.embedding_model and self.settings.curriculum_vector_enabled:
                state.embedding_model = self.settings.curriculum_embedding_model
            state.updated_at = state.updated_at or timestamp
        stale_sources = (
            self.db.query(CurriculumSourceModel)
            .filter(~CurriculumSourceModel.source.in_(list(grouped) or [""]))
            .all()
        )
        for state in stale_sources:
            self.db.delete(state)
        self.db.flush()

    def delete_source(self, source: str) -> int:
        normalized_source = normalize_source(source)
        try:
            self.vector_store.delete_source(normalized_source)
        except Exception:
            pass
        deleted = (
            self.db.query(CurriculumChunkModel)
            .filter(CurriculumChunkModel.source == normalized_source)
            .delete(synchronize_session=False)
        )
        self.db.query(CurriculumSourceModel).filter(
            CurriculumSourceModel.source == normalized_source
        ).delete(synchronize_session=False)
        return int(deleted)

    def rebuild_vector_index(self) -> int:
        chunks = (
            self.db.query(CurriculumChunkModel)
            .order_by(CurriculumChunkModel.source, CurriculumChunkModel.source_index)
            .all()
        )
        timestamp = now_iso()
        try:
            count = self.vector_store.rebuild(chunks)
        except Exception as exc:
            error = str(exc)
            self.ensure_source_records()
            for state in self.db.query(CurriculumSourceModel).all():
                state.vector_status = "error"
                state.vector_chunk_count = 0
                state.embedding_model = self.settings.curriculum_embedding_model
                state.last_error = error
                state.updated_at = timestamp
            self.db.flush()
            raise CurriculumVectorUnavailable(error) from exc

        grouped_counts: Counter[str] = Counter(row.source for row in chunks)
        self.ensure_source_records()
        for state in self.db.query(CurriculumSourceModel).all():
            state.vector_status = "ready"
            state.vector_chunk_count = grouped_counts.get(state.source, 0)
            state.embedding_model = self.settings.curriculum_embedding_model
            state.last_error = ""
            state.updated_at = timestamp
        self.db.flush()
        try:
            self.vector_store.snapshot()
        except Exception:
            pass
        return count

    def status(self) -> dict:
        self.ensure_source_records()
        vector_status = self.vector_store.status()
        vector_status.update(
            {
                "database_chunk_count": self.db.query(CurriculumChunkModel).count(),
                "source_count": self.db.query(CurriculumSourceModel).count(),
                "candidate_k": self.settings.curriculum_candidate_k,
                "top_k": self.settings.curriculum_top_k,
                "vector_min_similarity": self.settings.curriculum_vector_min_similarity,
                "vector_weight": self.settings.curriculum_hybrid_vector_weight,
                "bm25_weight": self.settings.curriculum_hybrid_bm25_weight,
            }
        )
        return vector_status

    def retrieve(self, query: str, top_k: int | None = None) -> list[CurriculumSearchResult]:
        query = query.strip()
        if not query:
            return []

        chunks = (
            self.db.query(CurriculumChunkModel)
            .order_by(CurriculumChunkModel.source, CurriculumChunkModel.source_index)
            .all()
        )
        if not chunks:
            return []

        bm25_raw = bm25_scores(query, chunks)
        bm25_maximum = max(bm25_raw.values(), default=0.0) or 1.0
        bm25_normalized = {
            chunk_id: score / bm25_maximum for chunk_id, score in bm25_raw.items()
        }
        vector_raw: dict[int, float] = {}
        vector_error = self.vector_store.error
        vector_used = False
        if self.settings.curriculum_vector_enabled and self.vector_store.available:
            try:
                vector_hits = self.vector_store.query(
                    query,
                    self.settings.curriculum_candidate_k,
                )
                valid_ids = {chunk.id for chunk in chunks}
                vector_raw = {
                    hit.chunk_id: hit.score
                    for hit in vector_hits
                    if hit.chunk_id in valid_ids
                    and hit.score >= self.settings.curriculum_vector_min_similarity
                }
                vector_used = True
                vector_error = ""
            except Exception as exc:
                vector_error = str(exc)
        vector_maximum = max(vector_raw.values(), default=0.0) or 1.0
        vector_normalized = {
            chunk_id: score / vector_maximum for chunk_id, score in vector_raw.items()
        }
        # 向量检索成功时，以余弦阈值作为最终候选门槛；BM25 参与排序，
        # 但不能让低于阈值的片段重新进入候选集。向量不可用时才纯 BM25 降级。
        candidate_ids = (
            set(vector_normalized) if vector_used else set(bm25_normalized)
        )
        if not candidate_ids:
            self.last_retrieval = {
                "mode": "local_hybrid_empty" if vector_used else "local_bm25_empty",
                "vector_error": vector_error,
            }
            return []

        query_levels = detect_education_levels(query)
        query_subjects = detect_subjects(query)
        ranked: list[tuple[CurriculumChunkModel, float, float, float, float]] = []
        has_vector_results = bool(vector_normalized)
        has_bm25_results = bool(bm25_normalized)
        vector_weight = self.settings.curriculum_hybrid_vector_weight if has_vector_results else 0.0
        bm25_weight = self.settings.curriculum_hybrid_bm25_weight if has_bm25_results else 0.0
        total_weight = vector_weight + bm25_weight or 1.0
        for chunk in chunks:
            if chunk.id not in candidate_ids:
                continue
            chunk_levels = detect_education_levels(f"{chunk.source}\n{chunk.content}")
            if query_levels and chunk_levels and query_levels.isdisjoint(chunk_levels):
                continue
            source_subjects = detect_subjects(chunk.source)
            if query_subjects and source_subjects and query_subjects.isdisjoint(source_subjects):
                continue
            # 对外展示原始分；相对归一化分只用于融合排序。
            bm25_score = bm25_raw.get(chunk.id, 0.0)
            bm25_relative_score = bm25_normalized.get(chunk.id, 0.0)
            vector_score = vector_raw.get(chunk.id, 0.0)
            vector_relative_score = vector_normalized.get(chunk.id, 0.0)
            if vector_used:
                fusion_score = (
                    vector_relative_score * vector_weight
                    + bm25_relative_score * bm25_weight
                ) / total_weight
                score = fusion_score
                if self.settings.curriculum_rerank_enabled:
                    score = (
                        fusion_score * 0.75
                        + query_token_coverage(query, chunk.content) * 0.15
                        + phrase_score(query, chunk.content) * 0.10
                    )
            else:
                fusion_score = bm25_relative_score
                score = bm25_relative_score
                if self.settings.curriculum_rerank_enabled:
                    score = (
                        bm25_relative_score * 0.65
                        + query_token_coverage(query, chunk.content) * 0.25
                        + phrase_score(query, chunk.content) * 0.10
                    )
            ranked.append((chunk, score, bm25_score, vector_score, fusion_score))

        ranked.sort(key=lambda item: (-item[1], item[0].source, item[0].source_index))
        candidates = ranked[: self.settings.curriculum_candidate_k]
        limit = max(1, top_k or self.settings.curriculum_top_k)
        selected = select_diverse_scored_anchors(candidates, limit)
        chunk_map = {(chunk.source, chunk.source_index): chunk for chunk in chunks}

        mode = "local_hybrid" if vector_used else (
            "local_bm25_fallback"
            if self.settings.curriculum_vector_enabled
            else "local_bm25"
        )
        self.last_retrieval = {"mode": mode, "vector_error": vector_error}

        results: list[CurriculumSearchResult] = []
        for anchor, score, bm25_score, vector_score, fusion_score in selected:
            neighbors = [
                chunk_map.get((anchor.source, index))
                for index in range(anchor.source_index - 1, anchor.source_index + 2)
            ]
            neighbors = [item for item in neighbors if item is not None]
            results.append(
                CurriculumSearchResult(
                    chunk_id=anchor.id,
                    source=anchor.source,
                    source_index=anchor.source_index,
                    content="\n\n".join(item.content for item in neighbors),
                    score=round(score, 6),
                    chunk_ids=tuple(item.id for item in neighbors),
                    bm25_score=round(bm25_score, 6),
                    vector_score=round(vector_score, 6),
                    fusion_score=round(fusion_score, 6),
                    retrieval_mode=mode,
                )
            )
        return results

    @staticmethod
    def extract_text(path: Path) -> str:
        return CurriculumKnowledgeService.extract_bytes(path.name, path.read_bytes())

    @staticmethod
    def safe_source_name(filename: str) -> str:
        source = Path((filename or "").replace("\\", "/")).name.strip()
        if not source:
            raise CurriculumFileError("课标文件名不能为空")
        return source[:255]

    @staticmethod
    def extract_bytes(filename: str, data: bytes) -> str:
        source = CurriculumKnowledgeService.safe_source_name(filename)
        extension = Path(source).suffix.lower()
        if extension not in SUPPORTED_CURRICULUM_EXTENSIONS:
            supported = ", ".join(sorted(SUPPORTED_CURRICULUM_EXTENSIONS))
            raise CurriculumFileError(f"仅支持以下课标文件格式：{supported}")
        if not data:
            raise CurriculumFileError("不能上传空文件")
        try:
            if extension in {".md", ".txt"}:
                return extract_plain_bytes(data)
            if extension == ".pdf":
                return extract_pdf_bytes(data)
            if extension == ".docx":
                return extract_docx_bytes(data)
        except CurriculumFileError:
            raise
        except Exception as exc:
            raise CurriculumFileError(f"课标文件解析失败：{source}") from exc
        raise CurriculumFileError(f"不支持的课标文件格式：{extension}")


def now_iso() -> str:
    return dt.datetime.now(dt.timezone.utc).astimezone().isoformat()


def normalize_source(source: str) -> str:
    normalized = (source or "").replace("\\", "/").strip()
    if not normalized or normalized.startswith("/"):
        raise CurriculumFileError("课标来源名称不能为空")
    parts = [part for part in normalized.split("/") if part]
    if not parts or any(part in {".", ".."} for part in parts):
        raise CurriculumFileError("课标来源路径无效")
    return "/".join(parts)[:255]


def checksum_chunks(chunks: list[str]) -> str:
    payload = "\n\n".join(chunks).encode("utf-8")
    return hashlib.sha256(payload).hexdigest()


def chunk_text(content: str, *, size: int, overlap: int) -> list[str]:
    text = normalize_text(content)
    if not text:
        return []

    chunks: list[str] = []
    start = 0
    length = len(text)
    boundary_chars = "\n。！？；.!?;"
    minimum_boundary = max(1, int(size * 0.6))

    while start < length:
        target = min(length, start + size)
        end = target
        if target < length:
            window = text[start + minimum_boundary : target]
            boundary = max((window.rfind(char) for char in boundary_chars), default=-1)
            if boundary >= 0:
                end = start + minimum_boundary + boundary + 1

        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= length:
            break
        start = max(start + 1, end - overlap)

    return chunks


def bm25_scores(query: str, chunks: list[CurriculumChunkModel]) -> dict[int, float]:
    query_terms = Counter(tokenize(query))
    if not query_terms:
        return {}

    documents: list[tuple[int, Counter[str], int]] = []
    document_frequencies: Counter[str] = Counter()
    for chunk in chunks:
        token_counts = Counter(tokenize(chunk.content))
        documents.append((chunk.id, token_counts, sum(token_counts.values())))
        document_frequencies.update(token_counts.keys())

    average_length = sum(item[2] for item in documents) / len(documents) or 1.0
    total_documents = len(documents)
    scores: dict[int, float] = {}
    k1 = 1.5
    b = 0.75

    for chunk_id, token_counts, document_length in documents:
        score = 0.0
        length_normalizer = k1 * (1.0 - b + b * document_length / average_length)
        for term, query_frequency in query_terms.items():
            term_frequency = token_counts.get(term, 0)
            if term_frequency == 0:
                continue
            document_frequency = document_frequencies[term]
            inverse_frequency = math.log(
                1.0 + (total_documents - document_frequency + 0.5) / (document_frequency + 0.5)
            )
            query_boost = 1.0 + math.log(query_frequency)
            score += (
                inverse_frequency
                * query_boost
                * (term_frequency * (k1 + 1.0))
                / (term_frequency + length_normalizer)
            )
        if score > 0:
            scores[chunk_id] = score
    return scores


def select_diverse_anchors(
    candidates: list[tuple[CurriculumChunkModel, float]],
    limit: int,
) -> list[tuple[CurriculumChunkModel, float]]:
    selected: list[tuple[CurriculumChunkModel, float]] = []
    deferred: list[tuple[CurriculumChunkModel, float]] = []
    for item in candidates:
        chunk = item[0]
        overlaps = any(
            chosen.source == chunk.source
            and abs(chosen.source_index - chunk.source_index) <= 1
            for chosen, _ in selected
        )
        if overlaps:
            deferred.append(item)
        else:
            selected.append(item)
        if len(selected) >= limit:
            return selected

    for item in deferred:
        selected.append(item)
        if len(selected) >= limit:
            break
    return selected


def select_diverse_scored_anchors(
    candidates: list[
        tuple[CurriculumChunkModel, float, float, float, float]
    ],
    limit: int,
) -> list[tuple[CurriculumChunkModel, float, float, float, float]]:
    selected: list[tuple[CurriculumChunkModel, float, float, float, float]] = []
    deferred: list[tuple[CurriculumChunkModel, float, float, float, float]] = []
    for item in candidates:
        chunk = item[0]
        overlaps = any(
            chosen.source == chunk.source
            and abs(chosen.source_index - chunk.source_index) <= 1
            for chosen, *_ in selected
        )
        if overlaps:
            deferred.append(item)
        else:
            selected.append(item)
        if len(selected) >= limit:
            return selected

    for item in deferred:
        selected.append(item)
        if len(selected) >= limit:
            break
    return selected


def tokenize(text: str) -> list[str]:
    raw_tokens = re.findall(r"[a-zA-Z0-9_]+|[\u4e00-\u9fff]", text.lower())
    tokens = [token for token in raw_tokens if token not in TOKEN_STOPWORDS]
    chinese = "".join(char for char in text.lower() if "\u4e00" <= char <= "\u9fff")
    tokens.extend(chinese[index : index + 2] for index in range(max(0, len(chinese) - 1)))
    return tokens


def detect_education_levels(text: str) -> set[str]:
    return {
        level
        for level, keywords in EDUCATION_LEVEL_KEYWORDS.items()
        if any(keyword in text for keyword in keywords)
    }


def detect_subjects(text: str) -> set[str]:
    normalized = text.lower()
    return {
        subject
        for subject, keywords in SUBJECT_KEYWORDS.items()
        if any(keyword in normalized for keyword in keywords)
    }


def query_token_coverage(query: str, content: str) -> float:
    query_tokens = set(tokenize(query))
    if not query_tokens:
        return 0.0
    return len(query_tokens & set(tokenize(content))) / len(query_tokens)


def phrase_score(query: str, content: str) -> float:
    query_parts = [
        compact_text(part)
        for part in re.split(r"[\n，。！？、；：,.!?;:]", query)
        if len(compact_text(part)) >= 2
    ]
    if not query_parts:
        return 0.0
    compact_content = compact_text(content)
    return sum(1 for part in query_parts if part in compact_content) / len(query_parts)


def compact_text(text: str) -> str:
    return re.sub(r"\s+", "", text.lower())


def extract_plain_text(path: Path) -> str:
    return extract_plain_bytes(path.read_bytes())


def extract_plain_bytes(data: bytes) -> str:
    for encoding in ("utf-8-sig", "gb18030"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise CurriculumFileError("TXT/MD 文件编码无法识别，请使用 UTF-8 或 GB18030")


def extract_pdf(path: Path) -> str:
    return extract_pdf_bytes(path.read_bytes())


def extract_pdf_bytes(data: bytes) -> str:
    reader = PdfReader(BytesIO(data))
    if reader.is_encrypted and reader.decrypt("") == 0:
        raise CurriculumFileError("PDF 已加密，无法读取")
    pages = [normalize_text(page.extract_text() or "") for page in reader.pages]
    pages = remove_repeated_page_edges(pages)
    return "\n\n".join(page for page in pages if page)


def remove_repeated_page_edges(pages: list[str]) -> list[str]:
    if len(pages) < 2:
        return pages

    page_lines = [[line.strip() for line in page.splitlines() if line.strip()] for page in pages]
    edge_counts: Counter[str] = Counter()
    for lines in page_lines:
        if lines:
            edge_counts.update({lines[0], lines[-1]})
    threshold = max(2, math.ceil(len(page_lines) * 0.5))
    repeated_edges = {line for line, count in edge_counts.items() if count >= threshold}

    cleaned_pages = []
    for lines in page_lines:
        if lines and lines[0] in repeated_edges:
            lines = lines[1:]
        if lines and lines[-1] in repeated_edges:
            lines = lines[:-1]
        cleaned_pages.append("\n".join(lines))
    return cleaned_pages


def extract_docx(path: Path) -> str:
    return extract_docx_bytes(path.read_bytes())


def extract_docx_bytes(data: bytes) -> str:
    document = Document(BytesIO(data))
    blocks = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                blocks.append(" | ".join(cells))
    return "\n".join(blocks)


def normalize_text(text: str) -> str:
    text = text.replace("\x00", "").replace("\r\n", "\n").replace("\r", "\n")
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.splitlines()]
    return re.sub(r"\n{3,}", "\n\n", "\n".join(lines)).strip()
