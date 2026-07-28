import re
from pathlib import Path

from docx import Document
from pypdf import PdfReader
from sqlalchemy.orm import Session

from app.core.config import get_settings
from app.db.models import SessionFileModel


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


class FileExtractionError(ValueError):
    pass


class SessionFileService:
    @staticmethod
    def safe_display_name(filename: str) -> str:
        name = Path((filename or "").replace("\\", "/")).name.strip()
        if not name:
            raise FileExtractionError("文件名不能为空")
        return name[:255]

    @staticmethod
    def extension_for(filename: str) -> str:
        extension = Path(filename).suffix.lower()
        if extension not in SUPPORTED_EXTENSIONS:
            supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
            raise FileExtractionError(f"仅支持以下文件格式：{supported}")
        return extension

    @staticmethod
    def relative_storage_path(session_id: str, file_id: str, extension: str) -> Path:
        return Path(session_id) / f"{file_id}{extension}"

    @staticmethod
    def absolute_storage_path(relative_path: str | Path) -> Path:
        root = get_settings().upload_dir.resolve()
        resolved = (root / relative_path).resolve()
        if not resolved.is_relative_to(root):
            raise FileExtractionError("文件存储路径无效")
        return resolved

    @staticmethod
    def write_file(relative_path: Path, data: bytes) -> None:
        path = SessionFileService.absolute_storage_path(relative_path)
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_bytes(data)

    @staticmethod
    def extract_text(relative_path: Path, extension: str) -> str:
        path = SessionFileService.absolute_storage_path(relative_path)
        if extension in {".txt", ".md"}:
            text = SessionFileService._extract_plain_text(path)
        elif extension == ".pdf":
            text = SessionFileService._extract_pdf(path)
        elif extension == ".docx":
            text = SessionFileService._extract_docx(path)
        else:
            raise FileExtractionError("不支持的文件格式")

        normalized = SessionFileService._normalize_text(text)
        if not normalized:
            raise FileExtractionError("文件中未提取到可用文字；扫描版 PDF 暂不支持 OCR")
        return normalized

    @staticmethod
    def delete_stored_file(relative_path: str) -> None:
        try:
            path = SessionFileService.absolute_storage_path(relative_path)
        except FileExtractionError:
            return
        if path.exists():
            path.unlink()
        root = get_settings().upload_dir.resolve()
        parent = path.parent
        if parent != root and parent.exists() and not any(parent.iterdir()):
            parent.rmdir()

    @staticmethod
    def delete_session_files(db: Session, session_id: str) -> None:
        files = (
            db.query(SessionFileModel)
            .filter(SessionFileModel.session_id == session_id)
            .all()
        )
        for item in files:
            SessionFileService.delete_stored_file(item.stored_path)
        db.query(SessionFileModel).filter(SessionFileModel.session_id == session_id).delete()

    @staticmethod
    def _extract_plain_text(path: Path) -> str:
        data = path.read_bytes()
        for encoding in ("utf-8-sig", "gb18030"):
            try:
                return data.decode(encoding)
            except UnicodeDecodeError:
                continue
        raise FileExtractionError("TXT/MD 文件编码无法识别，请使用 UTF-8 或 GB18030")

    @staticmethod
    def _extract_pdf(path: Path) -> str:
        try:
            reader = PdfReader(str(path))
            if reader.is_encrypted and reader.decrypt("") == 0:
                raise FileExtractionError("PDF 已加密，无法读取")
            pages = [page.extract_text() or "" for page in reader.pages]
            return "\n\n".join(pages)
        except FileExtractionError:
            raise
        except Exception as exc:
            raise FileExtractionError("PDF 解析失败，请检查文件是否损坏") from exc

    @staticmethod
    def _extract_docx(path: Path) -> str:
        try:
            document = Document(str(path))
            blocks = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        blocks.append(" | ".join(cells))
            return "\n".join(blocks)
        except Exception as exc:
            raise FileExtractionError("DOCX 解析失败，请检查文件是否损坏") from exc

    @staticmethod
    def _normalize_text(text: str) -> str:
        text = text.replace("\x00", "").replace("\r\n", "\n").replace("\r", "\n")
        lines = [line.rstrip() for line in text.splitlines()]
        return re.sub(r"\n{3,}", "\n\n", "\n".join(lines)).strip()
