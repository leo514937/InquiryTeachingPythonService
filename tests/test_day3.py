import json
import os
import shutil
import tempfile
import unittest
import uuid
import zipfile
from io import BytesIO
from pathlib import Path
from unittest.mock import patch

from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker


TEST_DIR = Path(tempfile.mkdtemp(prefix="inquiry-agent-architecture-"))
TEST_DB = TEST_DIR / "agents.db"
os.environ["DATABASE_URL"] = f"sqlite:///{TEST_DB.as_posix()}"
os.environ["UPLOAD_DIR"] = str(TEST_DIR / "uploads")
os.environ["LLM_API_KEY"] = ""
os.environ["DIFY_DATASET_API_KEY"] = ""
os.environ["DIFY_DATASET_ID"] = ""
os.environ["DIFY_STAGE_AGENT_MODE"] = "mock"
os.environ["DIFY_STAGE_AGENTS_JSON"] = ""
os.environ["CURRICULUM_VECTOR_ENABLED"] = "false"
os.environ["FRONTEND_ORIGIN"] = (
    "http://127.0.0.1:5173,"
    "http://localhost:5173,"
    "http://152.136.39.252:5173"
)

from fastapi.testclient import TestClient
from docx import Document
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

from app.core.config import get_settings
from app.db.database import Base, SessionLocal, engine
from app.db.models import (
    AgentConversationModel,
    AppSettingModel,
    ChatTurnModel,
    CurriculumChunkModel,
    CurriculumSourceModel,
    DraftProposalModel,
    MessageModel,
    RagRecordModel,
    SessionFileModel,
    SessionModel,
)
from app.main import app
from app.services.auth_service import register_user
from app.services.chat_interrupt_service import chat_interruptions
from app.services.context_service import ContextService
from app.services.curriculum_knowledge_service import (
    CurriculumKnowledgeService,
    bm25_scores,
    detect_subjects,
)
from app.services.curriculum_vector_service import CurriculumVectorHit
from app.services.prompt_service import PromptService
from app.services.session_file_service import SessionFileService
from app.workflow.flows import get_flow


EXPECTED_STAGE_AGENTS = [
    "stage_observation_start",
    "stage_question_refine",
    "stage_hypothesis",
    "stage_experiment_design",
    "stage_new_questions",
    "stage_conclusion",
    "stage_extension",
]

EXPECTED_INSECT_HOTEL_STAGES = [
    ("natural_materials", "stage_observation_start"),
    ("habitat_needs", "stage_question_refine"),
    ("structure_design", "stage_hypothesis"),
    ("build_and_sensing", "stage_experiment_design"),
    ("settlement_observation", "stage_conclusion"),
    ("iteration_sharing", "stage_extension"),
]


def make_text_pdf(text: str) -> bytes:
    output = BytesIO()
    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    font_ref = writer._add_object(font)
    page[NameObject("/Resources")] = DictionaryObject(
        {NameObject("/Font"): DictionaryObject({NameObject("/F1"): font_ref})}
    )
    stream = DecodedStreamObject()
    stream.set_data(f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode("ascii"))
    page[NameObject("/Contents")] = writer._add_object(stream)
    writer.write(output)
    return output.getvalue()


def make_docx() -> bytes:
    output = BytesIO()
    document = Document()
    document.add_paragraph("DOCX_PARAGRAPH_MARKER")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "DOCX_TABLE_LEFT"
    table.cell(0, 1).text = "DOCX_TABLE_RIGHT"
    document.save(output)
    return output.getvalue()


def parse_sse(text: str) -> list[tuple[str, dict]]:
    events = []
    for block in text.strip().split("\n\n"):
        event_name = "message"
        data_lines = []
        for line in block.splitlines():
            if line.startswith("event:"):
                event_name = line[6:].strip()
            elif line.startswith("data:"):
                data_lines.append(line[5:].strip())
        if data_lines:
            events.append((event_name, json.loads("\n".join(data_lines))))
    return events


class FakeCurriculumVectorStore:
    def __init__(self, *, available: bool = True, error: str = ""):
        self.available = available
        self.error = error
        self.rows = []
        self.deleted_sources = []

    def index_chunks(self, chunks):
        if not self.available:
            raise RuntimeError(self.error or "fake vector unavailable")
        self.rows.extend(chunks)
        return len(chunks)

    def delete_source(self, source):
        self.deleted_sources.append(source)
        self.rows = [row for row in self.rows if row.source != source]

    def query(self, _query, top_k):
        if not self.available:
            raise RuntimeError(self.error or "fake vector unavailable")
        return [
            CurriculumVectorHit(chunk_id=int(row.id), score=0.95 - index * 0.05)
            for index, row in enumerate(self.rows[:top_k])
        ]

    def rebuild(self, chunks):
        if not self.available:
            raise RuntimeError(self.error or "fake vector unavailable")
        self.rows = list(chunks)
        return len(chunks)

    def snapshot(self):
        return None

    def status(self):
        return {
            "enabled": True,
            "required": False,
            "available": self.available,
            "dependency_ready": self.available,
            "model": "fake-local-model",
            "model_dir": "fake",
            "device": "cpu",
            "vector_dir": "fake",
            "collection": "curriculum_chunks",
            "vector_count": len(self.rows),
            "rebuild_required": False,
            "error": self.error,
        }


class AgentArchitectureApiTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.client = TestClient(app)
        cls.username = f"test_{uuid.uuid4().hex[:10]}"
        response = cls.client.post(
            "/api/auth/register",
            json={"username": cls.username, "password": "test-password-123"},
        )
        if response.status_code != 200:
            raise AssertionError(response.text)

    @classmethod
    def tearDownClass(cls):
        cls.client.close()
        engine.dispose()
        shutil.rmtree(TEST_DIR, ignore_errors=True)

    def create_session(
        self,
        flow_name: str = "inquiry_7_stage",
        topic: str = "光的折射",
    ) -> tuple[str, str]:
        response = self.client.post(
            "/api/sessions",
            json={"topic": topic, "flow_name": flow_name},
        )
        self.assertEqual(response.status_code, 200)
        data = response.json()["data"]
        return data["id"], data["current_stage"]["id"]

    def delete_session(self, session_id: str) -> None:
        response = self.client.delete(f"/api/sessions/{session_id}")
        self.assertEqual(response.status_code, 200)

    def stream_chat(self, session_id: str, message: str):
        response = self.client.post(
            f"/api/sessions/{session_id}/chat",
            json={"type": "chat", "message": message},
        )
        self.assertEqual(response.status_code, 200)
        return response, parse_sse(response.text)

    def get_session(self, session_id: str) -> dict:
        response = self.client.get(f"/api/sessions/{session_id}")
        self.assertEqual(response.status_code, 200)
        return response.json()["data"]

    def get_messages(self, session_id: str) -> list[dict]:
        response = self.client.get(f"/api/sessions/{session_id}/messages")
        self.assertEqual(response.status_code, 200)
        return response.json()["data"]

    def upload_file(
        self,
        session_id: str,
        name: str,
        content: bytes,
        content_type: str,
    ):
        return self.client.post(
            f"/api/sessions/{session_id}/files",
            files={"file": (name, content, content_type)},
        )

    def get_files(self, session_id: str) -> list[dict]:
        response = self.client.get(f"/api/sessions/{session_id}/files")
        self.assertEqual(response.status_code, 200)
        return response.json()["data"]

    def set_draft_mode(self, session_id: str, enabled: bool) -> dict:
        response = self.client.put(
            f"/api/sessions/{session_id}/draft-mode",
            json={"enabled": enabled},
        )
        self.assertEqual(response.status_code, 200)
        return response.json()["data"]

    def test_seven_stages_bind_to_seven_stage_agents(self):
        flow_response = self.client.get("/api/flows")
        self.assertEqual(flow_response.status_code, 200)
        inquiry_flow = next(
            item
            for item in flow_response.json()["data"]
            if item["name"] == "inquiry_7_stage"
        )
        self.assertEqual(
            [stage["agent_id"] for stage in inquiry_flow["stages"]],
            EXPECTED_STAGE_AGENTS,
        )

        session_id, _ = self.create_session()
        try:
            agent_response = self.client.get(f"/api/sessions/{session_id}/dify_agents")
            self.assertEqual(agent_response.status_code, 200)
            agents = agent_response.json()["data"]
            self.assertEqual([item["id"] for item in agents], EXPECTED_STAGE_AGENTS)
            self.assertTrue(all("stage_id" in item for item in agents))
            self.assertTrue(all(item["configured"] for item in agents))
            self.assertTrue(all(item["mode"] == "prompt" for item in agents))
        finally:
            self.delete_session(session_id)

    def test_insect_hotel_flow_is_listed_with_expected_stages(self):
        flow_response = self.client.get("/api/flows")
        self.assertEqual(flow_response.status_code, 200)

        insect_flow = next(
            item
            for item in flow_response.json()["data"]
            if item["name"] == "insect_hotel_project"
        )

        self.assertEqual(insect_flow["display_name"], "昆虫旅馆项目探究流")
        self.assertEqual(insect_flow["stage_count"], 6)
        self.assertEqual(
            [(stage["id"], stage["agent_id"]) for stage in insect_flow["stages"]],
            EXPECTED_INSECT_HOTEL_STAGES,
        )

    def test_insect_hotel_session_initializes_expected_outputs_and_agents(self):
        session_id, stage_id = self.create_session(
            flow_name="insect_hotel_project",
            topic="昆虫旅馆",
        )
        try:
            session = self.get_session(session_id)
            self.assertEqual(session["flow_name"], "insect_hotel_project")
            self.assertEqual(session["flow_display_name"], "昆虫旅馆项目探究流")
            self.assertEqual(session["current_stage"]["id"], "natural_materials")
            self.assertEqual(stage_id, "natural_materials")
            self.assertEqual(len(session["outputs"]), 6)
            self.assertEqual(
                [item["stage_id"] for item in session["outputs"]],
                [stage_id for stage_id, _ in EXPECTED_INSECT_HOTEL_STAGES],
            )

            agent_response = self.client.get(f"/api/sessions/{session_id}/dify_agents")
            self.assertEqual(agent_response.status_code, 200)
            agents = agent_response.json()["data"]
            self.assertEqual(
                [item["id"] for item in agents],
                [agent_id for _, agent_id in EXPECTED_INSECT_HOTEL_STAGES],
            )
            self.assertNotIn("stage_new_questions", [item["id"] for item in agents])
        finally:
            self.delete_session(session_id)

    def test_insect_hotel_flow_supports_stage_progression_and_stage_back(self):
        session_id, stage_id = self.create_session(
            flow_name="insect_hotel_project",
            topic="昆虫旅馆",
        )
        try:
            advance_response = self.client.post(
                f"/api/sessions/{session_id}/chat",
                json={
                    "type": "sys_action",
                    "action": "next_stage",
                    "final_content": "自然取材阶段定稿",
                },
            )
            self.assertEqual(advance_response.status_code, 200)

            session = self.get_session(session_id)
            self.assertEqual(session["current_stage_index"], 1)
            self.assertEqual(session["current_stage"]["id"], "habitat_needs")
            first_output = next(item for item in session["outputs"] if item["stage_id"] == stage_id)
            self.assertTrue(first_output["confirmed"])
            self.assertEqual(first_output["final_content"], "自然取材阶段定稿")

            rollback_response = self.client.post(
                f"/api/sessions/{session_id}/rollback",
                json={"steps": 1, "stage_back": True},
            )
            self.assertEqual(rollback_response.status_code, 200)
            rolled_session = rollback_response.json()["data"]["session"]
            self.assertEqual(rolled_session["current_stage_index"], 0)
            self.assertEqual(rolled_session["current_stage"]["id"], "natural_materials")
            rolled_output = next(item for item in rolled_session["outputs"] if item["stage_id"] == stage_id)
            self.assertFalse(rolled_output["confirmed"])
            self.assertEqual(rolled_output["final_content"], "")
        finally:
            self.delete_session(session_id)

    def test_auth_registration_validation_and_logout(self):
        username = f"auth_{uuid.uuid4().hex[:10]}"
        client = TestClient(app)
        try:
            invalid = client.post(
                "/api/auth/register",
                json={"username": "ab", "password": "short"},
            )
            self.assertEqual(invalid.status_code, 422)

            registered = client.post(
                "/api/auth/register",
                json={"username": username, "password": "valid-password-123"},
            )
            self.assertEqual(registered.status_code, 200)
            self.assertEqual(registered.json()["data"]["username"], username)
            self.assertFalse(registered.json()["data"]["is_admin"])
            self.assertEqual(client.get("/api/auth/me").status_code, 200)

            duplicate = client.post(
                "/api/auth/register",
                json={"username": username.upper(), "password": "valid-password-123"},
            )
            self.assertEqual(duplicate.status_code, 409)

            logout = client.post("/api/auth/logout")
            self.assertEqual(logout.status_code, 200)
            self.assertEqual(client.get("/api/auth/me").status_code, 401)

            invalid_login = client.post(
                "/api/auth/login",
                json={"username": username, "password": "wrong-password"},
            )
            self.assertEqual(invalid_login.status_code, 401)
            valid_login = client.post(
                "/api/auth/login",
                json={"username": username.upper(), "password": "valid-password-123"},
            )
            self.assertEqual(valid_login.status_code, 200)
            self.assertEqual(client.get("/api/auth/me").status_code, 200)
            for origin in (
                "http://127.0.0.1:5173",
                "http://localhost:5173",
                "http://152.136.39.252:5173",
            ):
                cors_response = client.options(
                    "/api/auth/register",
                    headers={
                        "Origin": origin,
                        "Access-Control-Request-Method": "POST",
                        "Access-Control-Request-Headers": "content-type",
                    },
                )
                self.assertEqual(cors_response.status_code, 200)
                self.assertEqual(
                    cors_response.headers.get("access-control-allow-origin"),
                    origin,
                )
                self.assertEqual(
                    cors_response.headers.get("access-control-allow-credentials"),
                    "true",
                )
        finally:
            client.close()

    def test_admin_channels_and_curriculum_file_permissions(self):
        ordinary_client = TestClient(app)
        admin_client = TestClient(app)
        disabled_client = TestClient(app)
        ordinary_username = f"ordinary_{uuid.uuid4().hex[:8]}"
        admin_username = f"admin_{uuid.uuid4().hex[:8]}"
        password = "valid-password-123"
        sources = [
            f"小学课标_{uuid.uuid4().hex}.md",
            f"课标表格_{uuid.uuid4().hex}.docx",
            f"课标文档_{uuid.uuid4().hex}.pdf",
        ]
        try:
            ordinary_registered = ordinary_client.post(
                "/api/auth/register",
                json={"username": ordinary_username, "password": password},
            )
            self.assertEqual(ordinary_registered.status_code, 200)
            self.assertFalse(ordinary_registered.json()["data"]["is_admin"])

            admin_registered = admin_client.post(
                "/api/auth/admin/register",
                json={"username": admin_username, "password": password},
            )
            self.assertEqual(admin_registered.status_code, 200)
            self.assertTrue(admin_registered.json()["data"]["is_admin"])

            admin_client.post("/api/auth/logout")
            self.assertEqual(
                admin_client.post(
                    "/api/auth/login",
                    json={"username": admin_username, "password": password},
                ).status_code,
                401,
            )
            self.assertEqual(
                admin_client.post(
                    "/api/auth/admin/login",
                    json={"username": admin_username, "password": password},
                ).status_code,
                200,
            )

            ordinary_client.post("/api/auth/logout")
            self.assertEqual(
                ordinary_client.post(
                    "/api/auth/admin/login",
                    json={"username": ordinary_username, "password": password},
                ).status_code,
                401,
            )
            self.assertEqual(
                ordinary_client.post(
                    "/api/auth/login",
                    json={"username": ordinary_username, "password": password},
                ).status_code,
                200,
            )

            with patch.object(get_settings(), "admin_registration_enabled", False):
                disabled = disabled_client.post(
                    "/api/auth/admin/register",
                    json={
                        "username": f"disabled_{uuid.uuid4().hex[:8]}",
                        "password": password,
                    },
                )
            self.assertEqual(disabled.status_code, 403)

            self.assertEqual(ordinary_client.get("/api/curriculum/files").status_code, 200)
            denied_upload = ordinary_client.post(
                "/api/curriculum/files",
                files={"file": (sources[0], b"ordinary cannot upload", "text/markdown")},
            )
            self.assertEqual(denied_upload.status_code, 403)

            uploads = [
                (sources[0], "小学三年级人工智能活动应使用图形化工具。".encode("utf-8"), "text/markdown"),
                (
                    sources[1],
                    make_docx(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
                (sources[2], make_text_pdf("CURRICULUM_API_PDF"), "application/pdf"),
            ]
            for name, data, content_type in uploads:
                uploaded = admin_client.post(
                    "/api/curriculum/files",
                    files={"file": (name, data, content_type)},
                )
                self.assertEqual(uploaded.status_code, 200, uploaded.text)
                self.assertEqual(uploaded.json()["data"]["source"], name)
                self.assertGreater(uploaded.json()["data"]["chunk_count"], 0)

            replaced = admin_client.post(
                "/api/curriculum/files",
                files={
                    "file": (
                        sources[0],
                        "小学三年级活动需要过程评价与安全边界。".encode("utf-8"),
                        "text/markdown",
                    )
                },
            )
            self.assertEqual(replaced.status_code, 200)
            with SessionLocal() as db:
                markdown_chunks = (
                    db.query(CurriculumChunkModel)
                    .filter(CurriculumChunkModel.source == sources[0])
                    .all()
                )
                self.assertTrue(markdown_chunks)
                self.assertTrue(all("图形化工具" not in row.content for row in markdown_chunks))

            failed_replacement = admin_client.post(
                "/api/curriculum/files",
                files={"file": (sources[1], b"not-a-docx", "application/octet-stream")},
            )
            self.assertEqual(failed_replacement.status_code, 400)
            with SessionLocal() as db:
                self.assertGreater(
                    db.query(CurriculumChunkModel)
                    .filter(CurriculumChunkModel.source == sources[1])
                    .count(),
                    0,
                )

            ordinary_sources = {
                item["source"]
                for item in ordinary_client.get("/api/curriculum/files").json()["data"]
            }
            self.assertTrue(set(sources).issubset(ordinary_sources))
            self.assertEqual(
                ordinary_client.delete(
                    "/api/curriculum/files",
                    params={"source": sources[0]},
                ).status_code,
                403,
            )

            for source in sources:
                deleted = admin_client.delete(
                    "/api/curriculum/files",
                    params={"source": source},
                )
                self.assertEqual(deleted.status_code, 200)
        finally:
            with SessionLocal() as db:
                db.query(CurriculumChunkModel).filter(
                    CurriculumChunkModel.source.in_(sources)
                ).delete(synchronize_session=False)
                db.commit()
            ordinary_client.close()
            admin_client.close()
            disabled_client.close()

    def test_first_registration_claims_legacy_sessions_and_chat_mode(self):
        database_path = TEST_DIR / f"legacy-{uuid.uuid4().hex}.db"
        local_engine = create_engine(
            f"sqlite:///{database_path.as_posix()}",
            connect_args={"check_same_thread": False},
        )
        Base.metadata.create_all(local_engine)
        LocalSession = sessionmaker(autocommit=False, autoflush=False, bind=local_engine)
        try:
            with LocalSession() as db:
                db.add(
                    AppSettingModel(
                        key="global_chat_mode",
                        value="subagent",
                        updated_at="2026-01-01T00:00:00+00:00",
                    )
                )
                legacy = SessionModel(
                    id="legacy_session",
                    owner_user_id=None,
                    title="旧会话",
                    topic="旧课题",
                    flow_name="inquiry_7_stage",
                    current_stage_index=0,
                    status="active",
                    draft_mode_enabled=0,
                    created_at="2026-01-01T00:00:00+00:00",
                    updated_at="2026-01-01T00:00:00+00:00",
                )
                db.add(legacy)
                db.commit()

                user, _ = register_user(db, "legacy_owner", "valid-password-123")
                db.refresh(legacy)

                self.assertEqual(legacy.owner_user_id, user.id)
                self.assertEqual(user.chat_mode, "subagent")
        finally:
            local_engine.dispose()

    def test_users_have_isolated_sessions_and_chat_modes(self):
        alice = TestClient(app)
        bob = TestClient(app)
        try:
            for client, username in (
                (alice, f"alice_{uuid.uuid4().hex[:10]}"),
                (bob, f"bob_{uuid.uuid4().hex[:10]}"),
            ):
                response = client.post(
                    "/api/auth/register",
                    json={"username": username, "password": "valid-password-123"},
                )
                self.assertEqual(response.status_code, 200)

            alice_session = alice.post(
                "/api/sessions",
                json={"topic": "Alice 的课题", "flow_name": "inquiry_7_stage"},
            ).json()["data"]["id"]
            bob_session = bob.post(
                "/api/sessions",
                json={"topic": "Bob 的课题", "flow_name": "inquiry_7_stage"},
            ).json()["data"]["id"]

            self.assertNotEqual(alice_session, bob_session)
            self.assertEqual(
                [item["id"] for item in alice.get("/api/sessions").json()["data"]],
                [alice_session],
            )
            self.assertEqual(
                [item["id"] for item in bob.get("/api/sessions").json()["data"]],
                [bob_session],
            )
            self.assertEqual(alice.get(f"/api/sessions/{bob_session}").status_code, 404)
            self.assertEqual(bob.delete(f"/api/sessions/{alice_session}").status_code, 404)

            self.assertEqual(
                alice.put(
                    "/api/settings/chat-mode",
                    json={"chat_mode": "subagent"},
                ).status_code,
                200,
            )
            self.assertEqual(
                bob.get("/api/settings/chat-mode").json()["data"]["chat_mode"],
                "main",
            )
            alice_events = parse_sse(
                alice.post(
                    f"/api/sessions/{alice_session}/chat",
                    json={"type": "chat", "message": "Alice 的第一轮提问"},
                ).text
            )
            bob_events = parse_sse(
                bob.post(
                    f"/api/sessions/{bob_session}/chat",
                    json={"type": "chat", "message": "Bob 的第一轮提问"},
                ).text
            )
            self.assertEqual(alice_events[0][1]["chat_mode"], "subagent")
            self.assertEqual(bob_events[0][1]["chat_mode"], "main")
        finally:
            alice.close()
            bob.close()

    def test_same_topic_different_flows_are_separate_and_flow_switch_cannot_clear_data(self):
        first_response = self.client.post(
            "/api/sessions",
            json={"topic": "光的折射", "flow_name": "inquiry_7_stage"},
        )
        second_response = self.client.post(
            "/api/sessions",
            json={"topic": "光的折射", "flow_name": "three_step_inquiry"},
        )
        self.assertEqual(first_response.status_code, 200)
        self.assertEqual(second_response.status_code, 200)

        first = first_response.json()["data"]
        second = second_response.json()["data"]
        self.assertNotEqual(first["id"], second["id"])
        self.assertEqual(first["flow_name"], "inquiry_7_stage")
        self.assertEqual(second["flow_name"], "three_step_inquiry")
        self.assertNotEqual(len(first["outputs"]), len(second["outputs"]))
        listed_ids = {
            item["id"]
            for item in self.client.get("/api/sessions").json()["data"]
        }
        self.assertTrue({first["id"], second["id"]}.issubset(listed_ids))

        first_stage_id = first["outputs"][0]["stage_id"]
        with SessionLocal() as db:
            db.add(
                MessageModel(
                    id="msg_isolation_test",
                    session_id=first["id"],
                    stage_id=first_stage_id,
                    role="user",
                    content="只属于七阶段流程的消息",
                    agent_id=None,
                    message_type="chat",
                    created_at="2026-01-01T00:00:00+00:00",
                )
            )
            db.add(
                ChatTurnModel(
                    turn_id="turn_isolation_test",
                    session_id=first["id"],
                    stage_id=first_stage_id,
                    user_message_id="msg_isolation_test",
                    assistant_message_id="msg_isolation_test",
                    draft_before="",
                    draft_after="",
                    created_at="2026-01-01T00:00:00+00:00",
                )
            )
            db.add(
                DraftProposalModel(
                    id="proposal_isolation_test",
                    session_id=first["id"],
                    stage_id=first_stage_id,
                    base_content="旧草案",
                    candidate_content="新草案",
                    diff_json="[]",
                    status="pending",
                    created_at="2026-01-01T00:00:00+00:00",
                    updated_at="2026-01-01T00:00:00+00:00",
                )
            )
            db.add(
                RagRecordModel(
                    id="rag_isolation_test",
                    session_id=first["id"],
                    stage_id=first_stage_id,
                    query="隔离测试",
                    context="只属于第一个会话",
                    source_json="[]",
                    created_at="2026-01-01T00:00:00+00:00",
                )
            )
            db.add(
                AgentConversationModel(
                    id="conversation_isolation_test",
                    session_id=first["id"],
                    agent_id="stage_observation_start",
                    conversation_id="conversation-1",
                    created_at="2026-01-01T00:00:00+00:00",
                    updated_at="2026-01-01T00:00:00+00:00",
                )
            )
            db.commit()

        switched = self.client.post(
            f"/api/sessions/{first['id']}/select_flow",
            json={"flow_name": "three_step_inquiry", "clear_messages": True},
        )
        self.assertEqual(switched.status_code, 409)
        self.assertIn("流程不可修改", switched.json()["detail"])

        first_after = self.get_session(first["id"])
        second_after = self.get_session(second["id"])
        self.assertEqual(first_after["flow_name"], "inquiry_7_stage")
        self.assertEqual(len(first_after["outputs"]), 7)
        self.assertEqual(len(second_after["outputs"]), 3)
        self.assertEqual(len(self.get_messages(first["id"])), 1)
        self.assertEqual(self.get_messages(second["id"]), [])

        with SessionLocal() as db:
            self.assertEqual(
                db.query(ChatTurnModel)
                .filter(ChatTurnModel.session_id == first["id"])
                .count(),
                1,
            )
            self.assertEqual(
                db.query(DraftProposalModel)
                .filter(DraftProposalModel.session_id == first["id"])
                .count(),
                1,
            )
            self.assertEqual(
                db.query(RagRecordModel)
                .filter(RagRecordModel.session_id == first["id"])
                .count(),
                1,
            )
            self.assertEqual(
                db.query(AgentConversationModel)
                .filter(AgentConversationModel.session_id == first["id"])
                .count(),
                1,
            )

        self.delete_session(first["id"])
        self.assertEqual(self.get_session(second["id"])["flow_name"], "three_step_inquiry")
        self.delete_session(second["id"])

    def test_chat_stream_defaults_to_main_mode_and_persists_two_messages(self):
        session_id, stage_id = self.create_session()
        try:
            response, events = self.stream_chat(session_id, "用筷子折弯现象导入")
            event_names = [name for name, _ in events]

            self.assertEqual(event_names[0], "stage")
            agent_events = [
                (index, data)
                for index, (name, data) in enumerate(events)
                if name == "agent"
            ]
            self.assertEqual(
                [data["message_type"] for _, data in agent_events],
                ["main_tutor"],
            )
            main_delta_indices = [
                index
                for index, (name, data) in enumerate(events)
                if name == "delta" and data.get("message_type") == "main_tutor"
            ]
            self.assertTrue(main_delta_indices)
            self.assertNotIn("draft", event_names)
            self.assertEqual(event_names[-1], "done")
            self.assertFalse(events[-1][1]["degraded"])
            self.assertFalse(events[-1][1]["draft_mode_enabled"])
            self.assertFalse(events[-1][1]["draft_updated"])
            self.assertEqual(response.headers["cache-control"], "no-cache")
            self.assertEqual(response.headers["x-accel-buffering"], "no")

            messages = self.get_messages(session_id)
            self.assertEqual(
                [item["message_type"] for item in messages],
                ["chat", "main_tutor"],
            )
            self.assertEqual(
                [item["agent_id"] for item in messages],
                [None, "main_agent"],
            )
            self.assertNotIn("===DRAFT_START===", messages[1]["content"])

            session = self.get_session(session_id)
            output = next(item for item in session["outputs"] if item["stage_id"] == stage_id)
            self.assertFalse(output["draft_content"])

            with SessionLocal() as db:
                turn = (
                    db.query(ChatTurnModel)
                    .filter(ChatTurnModel.session_id == session_id)
                    .one()
                )
                self.assertFalse(turn.expert_message_id)
                self.assertTrue(turn.rag_record_id)
                self.assertEqual(
                    db.query(RagRecordModel)
                    .filter(RagRecordModel.session_id == session_id)
                    .count(),
                    1,
                )
                self.assertEqual(
                    db.query(AgentConversationModel)
                    .filter(AgentConversationModel.session_id == session_id)
                    .count(),
                    0,
                )
        finally:
            self.delete_session(session_id)

    def test_curriculum_files_ingest_idempotently_and_retrieve_with_bm25(self):
        source = f"小学信息科技课标_{uuid.uuid4().hex}.md"
        wrong_level_source = f"初中信息科技课标_{uuid.uuid4().hex}.md"
        curriculum_dir = TEST_DIR / f"curriculum_{uuid.uuid4().hex}"
        curriculum_dir.mkdir(parents=True, exist_ok=True)
        markdown_path = curriculum_dir / source
        markdown_path.write_text(
            "小学三年级学生以直观体验为主，人工智能活动应采用图形化工具，控制任务难度。",
            encoding="utf-8",
        )
        try:
            with SessionLocal() as db:
                service = CurriculumKnowledgeService(db)
                first_count = service.ingest_file(markdown_path, source=source)
                db.commit()
                self.assertGreater(first_count, 0)

                markdown_path.write_text(
                    "小学三年级学生适合图形化人工智能活动，并应设置清晰的安全边界与过程评价。",
                    encoding="utf-8",
                )
                second_count = service.ingest_file(markdown_path, source=source)
                db.commit()
                stored = (
                    db.query(CurriculumChunkModel)
                    .filter(CurriculumChunkModel.source == source)
                    .all()
                )
                self.assertEqual(len(stored), second_count)
                self.assertTrue(all("直观体验为主" not in item.content for item in stored))

                service.ingest(
                    wrong_level_source,
                    "初中七年级学生可以使用 Python 编程完成复杂的人工智能模型训练活动。",
                )
                db.commit()

                results = service.retrieve("小学三年级人工智能活动的安全与评价", top_k=2)
                self.assertTrue(results)
                self.assertEqual(results[0].source, source)
                self.assertIn("安全边界", results[0].content)
                self.assertNotIn(wrong_level_source, [item.source for item in results])
        finally:
            with SessionLocal() as db:
                db.query(CurriculumChunkModel).filter(
                    CurriculumChunkModel.source.in_([source, wrong_level_source])
                ).delete(synchronize_session=False)
                db.commit()
            shutil.rmtree(curriculum_dir, ignore_errors=True)

    def test_curriculum_parser_supports_txt_pdf_and_docx(self):
        curriculum_dir = TEST_DIR / f"curriculum_formats_{uuid.uuid4().hex}"
        curriculum_dir.mkdir(parents=True, exist_ok=True)
        text_path = curriculum_dir / "standard.txt"
        pdf_path = curriculum_dir / "standard.pdf"
        docx_path = curriculum_dir / "standard.docx"
        text_path.write_text("TXT_CURRICULUM_MARKER", encoding="utf-8")
        pdf_path.write_bytes(make_text_pdf("PDF_CURRICULUM_MARKER"))
        docx_path.write_bytes(make_docx())
        try:
            self.assertIn(
                "TXT_CURRICULUM_MARKER",
                CurriculumKnowledgeService.extract_text(text_path),
            )
            self.assertIn(
                "PDF_CURRICULUM_MARKER",
                CurriculumKnowledgeService.extract_text(pdf_path),
            )
            docx_text = CurriculumKnowledgeService.extract_text(docx_path)
            self.assertIn("DOCX_PARAGRAPH_MARKER", docx_text)
            self.assertIn("DOCX_TABLE_LEFT", docx_text)
        finally:
            shutil.rmtree(curriculum_dir, ignore_errors=True)

    def test_curriculum_hybrid_retrieval_and_vector_sync(self):
        source = f"初中物理课标_{uuid.uuid4().hex}.md"
        settings = get_settings()
        vector_store = FakeCurriculumVectorStore()
        try:
            with patch.object(settings, "curriculum_vector_enabled", True), patch.object(
                settings,
                "curriculum_embedding_model",
                "fake-local-model",
            ):
                with SessionLocal() as db:
                    service = CurriculumKnowledgeService(
                        db,
                        settings,
                        vector_store=vector_store,
                    )
                    count = service.ingest(
                        source,
                        "初中八年级学生通过推墙体验力的作用是相互的，并记录身体后退的证据。",
                    )
                    db.commit()
                    self.assertEqual(len(vector_store.rows), count)

                    query = "为什么推墙的人会向后移动"
                    results = service.retrieve(query, top_k=1)
                    self.assertTrue(results)
                    self.assertEqual(results[0].source, source)
                    self.assertEqual(results[0].vector_score, 0.95)
                    all_chunks = db.query(CurriculumChunkModel).all()
                    expected_bm25 = bm25_scores(query, all_chunks)[results[0].chunk_id]
                    self.assertAlmostEqual(results[0].bm25_score, expected_bm25, places=6)
                    self.assertEqual(results[0].retrieval_mode, "local_hybrid")

                    deleted = service.delete_source(source)
                    db.commit()
                    self.assertGreater(deleted, 0)
                    self.assertIn(source, vector_store.deleted_sources)
                    self.assertEqual(
                        db.query(CurriculumChunkModel)
                        .filter(CurriculumChunkModel.source == source)
                        .count(),
                        0,
                    )
        finally:
            with SessionLocal() as db:
                db.query(CurriculumSourceModel).filter(
                    CurriculumSourceModel.source == source
                ).delete(synchronize_session=False)
                db.query(CurriculumChunkModel).filter(
                    CurriculumChunkModel.source == source
                ).delete(synchronize_session=False)
                db.commit()

    def test_curriculum_subject_detection_only_filters_explicit_subjects(self):
        self.assertEqual(detect_subjects("八年级力的作用是相互的探究"), {"physics"})
        self.assertEqual(detect_subjects("初中函数与方程教学"), {"mathematics"})
        self.assertEqual(detect_subjects("适合八年级的探究活动"), set())

    def test_curriculum_vector_similarity_threshold(self):
        source = f"初中物理阈值测试_{uuid.uuid4().hex}.md"
        settings = get_settings()
        vector_store = FakeCurriculumVectorStore()
        try:
            with patch.object(settings, "curriculum_vector_enabled", True), patch.object(
                settings,
                "curriculum_vector_min_similarity",
                0.5,
            ):
                with SessionLocal() as db:
                    service = CurriculumKnowledgeService(
                        db,
                        settings,
                        vector_store=vector_store,
                    )
                    service.ingest(source, "初中八年级学生观察力的相互作用。")
                    db.commit()
                    chunk = db.query(CurriculumChunkModel).filter(
                        CurriculumChunkModel.source == source
                    ).one()

                    with patch.object(
                        vector_store,
                        "query",
                        return_value=[CurriculumVectorHit(chunk_id=chunk.id, score=0.499999)],
                    ):
                        self.assertEqual(service.retrieve("完全不同的检索词", top_k=1), [])

                    with patch.object(
                        vector_store,
                        "query",
                        return_value=[CurriculumVectorHit(chunk_id=chunk.id, score=0.5)],
                    ):
                        results = service.retrieve("完全不同的检索词", top_k=1)
                        self.assertEqual(len(results), 1)
                        self.assertEqual(results[0].vector_score, 0.5)
        finally:
            with SessionLocal() as db:
                db.query(CurriculumSourceModel).filter(
                    CurriculumSourceModel.source == source
                ).delete(synchronize_session=False)
                db.query(CurriculumChunkModel).filter(
                    CurriculumChunkModel.source == source
                ).delete(synchronize_session=False)
                db.commit()

    def test_curriculum_admin_export_import_and_permissions(self):
        source = f"export_{uuid.uuid4().hex}.md"
        admin_client = TestClient(app)
        ordinary_client = TestClient(app)
        try:
            password = "valid-password-123"
            admin_client.post(
                "/api/auth/admin/register",
                json={"username": f"export_admin_{uuid.uuid4().hex[:8]}", "password": password},
            )
            ordinary_client.post(
                "/api/auth/register",
                json={"username": f"export_user_{uuid.uuid4().hex[:8]}", "password": password},
            )
            uploaded = admin_client.post(
                "/api/curriculum/files",
                files={"file": (source, "初中实验应记录证据。".encode("utf-8"), "text/markdown")},
            )
            self.assertEqual(uploaded.status_code, 200, uploaded.text)
            self.assertEqual(ordinary_client.get("/api/curriculum/status").status_code, 403)
            self.assertEqual(ordinary_client.get("/api/curriculum/export").status_code, 403)
            self.assertEqual(ordinary_client.get("/api/curriculum/retrievals").status_code, 403)

            exported = admin_client.get("/api/curriculum/export")
            self.assertEqual(exported.status_code, 200)
            with zipfile.ZipFile(BytesIO(exported.content)) as archive:
                payload = json.loads(archive.read("curriculum.json").decode("utf-8"))
            exported_source = next(item for item in payload["sources"] if item["source"] == source)
            serialized = json.dumps(exported_source, ensure_ascii=False)
            self.assertIn("初中实验应记录证据", serialized)
            self.assertNotIn("password_hash", serialized)
            self.assertNotIn("auth_sessions", serialized)

            self.assertEqual(
                admin_client.delete("/api/curriculum/files", params={"source": source}).status_code,
                200,
            )
            imported = admin_client.post(
                "/api/curriculum/import",
                files={"file": ("curriculum-knowledge.zip", exported.content, "application/zip")},
            )
            self.assertEqual(imported.status_code, 200, imported.text)
            listed = admin_client.get("/api/curriculum/files").json()["data"]
            self.assertIn(source, {item["source"] for item in listed})
        finally:
            with SessionLocal() as db:
                CurriculumKnowledgeService(db).delete_source(source)
                db.commit()
            admin_client.close()
            ordinary_client.close()

    def test_chat_injects_curriculum_reference_and_persists_sources(self):
        source = f"小学信息科技课标_{uuid.uuid4().hex}.md"
        with SessionLocal() as db:
            CurriculumKnowledgeService(db).ingest(
                source,
                "小学三年级学生适合使用图形化工具体验人工智能，任务步骤应简短，并设置安全边界。",
            )
            db.commit()

        session_id, _ = self.create_session(topic="小学三年级人工智能")
        try:
            _, events = self.stream_chat(session_id, "设计一个适合小学三年级的人工智能活动")
            done = next(data for name, data in events if name == "done")
            self.assertEqual(done["rag_sources"], [source])
            self.assertTrue(done["rag_record_id"])

            messages = self.get_messages(session_id)
            self.assertIn(f"参考课标：{source}", messages[-1]["content"])

            with SessionLocal() as db:
                turn = (
                    db.query(ChatTurnModel)
                    .filter(ChatTurnModel.session_id == session_id)
                    .one()
                )
                record = db.query(RagRecordModel).filter(
                    RagRecordModel.id == turn.rag_record_id
                ).one()
                self.assertIn("<curriculum_reference>", record.context)
                metadata = json.loads(record.source_json)
                self.assertEqual(metadata["mode"], "local_bm25")
                self.assertEqual(metadata["records"][0]["source"], source)
        finally:
            self.delete_session(session_id)
            with SessionLocal() as db:
                db.query(CurriculumChunkModel).filter(
                    CurriculumChunkModel.source == source
                ).delete(synchronize_session=False)
                db.commit()

    def test_chat_cancel_endpoint_marks_active_stream_as_cancelled(self):
        session_id, _ = self.create_session()
        request_id = f"cancel_{uuid.uuid4().hex}"
        try:
            with SessionLocal() as db:
                user_id = (
                    db.query(SessionModel)
                    .filter(SessionModel.id == session_id)
                    .one()
                    .owner_user_id
                )
            chat_interruptions.register(request_id, session_id, user_id)

            response = self.client.post(
                f"/api/sessions/{session_id}/chat/{request_id}/cancel"
            )

            self.assertEqual(response.status_code, 200)
            self.assertTrue(response.json()["data"]["cancelled"])
            self.assertTrue(chat_interruptions.is_cancelled(request_id))
        finally:
            chat_interruptions.unregister(request_id)
            self.delete_session(session_id)

    def test_manual_draft_save(self):
        session_id, stage_id = self.create_session()
        try:
            content = "### 手动草稿\n教师已完成二次编辑。"
            response = self.client.put(
                f"/api/sessions/{session_id}/stages/{stage_id}/draft",
                json={"draft_content": content},
            )
            self.assertEqual(response.status_code, 200)
            self.assertEqual(response.json()["data"]["draft_content"], content)

            session = self.get_session(session_id)
            output = next(item for item in session["outputs"] if item["stage_id"] == stage_id)
            self.assertEqual(output["draft_content"], content)
        finally:
            self.delete_session(session_id)

    def test_first_draft_generation_writes_directly_without_pending_proposal(self):
        session_id, stage_id = self.create_session()
        try:
            self.set_draft_mode(session_id, True)
            _, events = self.stream_chat(session_id, "请先生成一版观察阶段草案")

            event_names = [name for name, _ in events]
            self.assertIn("draft", event_names)
            self.assertNotIn("proposal", event_names)

            done_payload = events[-1][1]
            self.assertTrue(done_payload["draft_updated"])
            self.assertIsNone(done_payload["draft_proposal"])
            self.assertEqual(done_payload["draft_request_kind"], "generate")
            self.assertEqual(done_payload["proposal_kind"], "generate")

            session = self.get_session(session_id)
            output = next(item for item in session["outputs"] if item["stage_id"] == stage_id)
            self.assertTrue(output["draft_content"])

            proposal_response = self.client.get(
                f"/api/sessions/{session_id}/draft-proposal",
                params={"stage_id": stage_id},
            )
            self.assertEqual(proposal_response.status_code, 200)
            self.assertIsNone(proposal_response.json()["data"])

            with SessionLocal() as db:
                self.assertEqual(
                    db.query(DraftProposalModel)
                    .filter(
                        DraftProposalModel.session_id == session_id,
                        DraftProposalModel.stage_id == stage_id,
                        DraftProposalModel.status == "pending",
                    )
                    .count(),
                    0,
                )
        finally:
            self.delete_session(session_id)

    def test_existing_draft_edit_still_creates_pending_proposal(self):
        session_id, stage_id = self.create_session()
        try:
            self.set_draft_mode(session_id, True)
            base_content = "### 观察阶段草案\n\n1. 学生先记录现象。\n2. 教师组织交流。"
            response = self.client.put(
                f"/api/sessions/{session_id}/stages/{stage_id}/draft",
                json={"draft_content": base_content},
            )
            self.assertEqual(response.status_code, 200)

            edit_response = self.client.post(
                f"/api/sessions/{session_id}/chat",
                json={
                    "type": "chat",
                    "message": "把第一条改得更具体一点",
                    "draft_request_kind": "edit",
                    "selection": {
                        "selected_text": "1. 学生先记录现象。",
                        "start_offset": 15,
                        "end_offset": 27,
                        "stage_id": stage_id,
                    },
                },
            )
            self.assertEqual(edit_response.status_code, 200)
            events = parse_sse(edit_response.text)
            event_names = [name for name, _ in events]

            self.assertIn("draft", event_names)
            self.assertIn("proposal", event_names)

            done_payload = events[-1][1]
            self.assertTrue(done_payload["draft_updated"])
            self.assertIsNotNone(done_payload["draft_proposal"])
            self.assertEqual(done_payload["draft_request_kind"], "edit")
            self.assertEqual(done_payload["proposal_kind"], "edit")

            proposal_response = self.client.get(
                f"/api/sessions/{session_id}/draft-proposal",
                params={"stage_id": stage_id},
            )
            self.assertEqual(proposal_response.status_code, 200)
            proposal_data = proposal_response.json()["data"]
            self.assertIsNotNone(proposal_data)
            self.assertEqual(proposal_data["proposal_kind"], "edit")
            with SessionLocal() as db:
                self.assertEqual(
                    db.query(DraftProposalModel)
                    .filter(
                        DraftProposalModel.session_id == session_id,
                        DraftProposalModel.stage_id == stage_id,
                        DraftProposalModel.status == "pending",
                    )
                    .count(),
                    1,
                )
        finally:
            self.delete_session(session_id)

    def test_rollback_removes_three_messages_and_restores_previous_draft(self):
        session_id, stage_id = self.create_session()
        try:
            self.set_draft_mode(session_id, True)
            self.stream_chat(session_id, "第一版观察任务")
            first_session = self.get_session(session_id)
            first_draft = next(
                item["draft_content"]
                for item in first_session["outputs"]
                if item["stage_id"] == stage_id
            )
            self.assertTrue(first_draft)

            second_response = self.client.post(
                f"/api/sessions/{session_id}/chat",
                json={
                    "type": "chat",
                    "message": "第二版证据记录任务",
                    "draft_request_kind": "generate",
                },
            )
            self.assertEqual(second_response.status_code, 200)
            second_events = parse_sse(second_response.text)
            proposal = second_events[-1][1]["draft_proposal"]
            self.assertIsNotNone(proposal)
            actions = [
                {"hunk_id": segment["id"], "action": "accept"}
                for segment in proposal["segments"]
                if segment["kind"] != "equal"
            ]
            apply_response = self.client.post(
                f"/api/sessions/{session_id}/draft-proposals/{proposal['id']}/actions",
                json={"actions": actions},
            )
            self.assertEqual(apply_response.status_code, 200)

            second_session = self.get_session(session_id)
            second_draft = next(
                item["draft_content"]
                for item in second_session["outputs"]
                if item["stage_id"] == stage_id
            )
            self.assertNotEqual(first_draft, second_draft)
            self.assertEqual(len(self.get_messages(session_id)), 4)

            response = self.client.post(
                f"/api/sessions/{session_id}/rollback",
                json={"steps": 1, "stage_back": False},
            )
            self.assertEqual(response.status_code, 200)
            rollback_data = response.json()["data"]
            self.assertEqual(len(rollback_data["deleted_message_ids"]), 2)
            self.assertEqual(rollback_data["restored_drafts"][stage_id], first_draft)
            self.assertEqual(len(self.get_messages(session_id)), 2)

            with SessionLocal() as db:
                self.assertEqual(
                    db.query(ChatTurnModel)
                    .filter(ChatTurnModel.session_id == session_id)
                    .count(),
                    1,
                )
        finally:
            self.delete_session(session_id)

    def test_conversation_ids_are_reused_per_stage_and_isolated_between_stages(self):
        session_id, _ = self.create_session()
        try:
            self.client.put("/api/settings/chat-mode", json={"chat_mode": "subagent"})
            self.stream_chat(session_id, "观察问题一")
            self.stream_chat(session_id, "观察问题二")
            with SessionLocal() as db:
                observation_row = (
                    db.query(AgentConversationModel)
                    .filter(
                        AgentConversationModel.session_id == session_id,
                        AgentConversationModel.agent_id == "stage_observation_start",
                    )
                    .one()
                )
                observation_conversation_id = observation_row.conversation_id
                self.assertEqual(
                    db.query(AgentConversationModel)
                    .filter(AgentConversationModel.session_id == session_id)
                    .count(),
                    1,
                )

            current = self.get_session(session_id)
            current_draft = current["outputs"][0]["draft_content"]
            advance = self.client.post(
                f"/api/sessions/{session_id}/chat",
                json={
                    "type": "sys_action",
                    "action": "next_stage",
                    "final_content": current_draft,
                },
            )
            self.assertEqual(advance.status_code, 200)
            self.stream_chat(session_id, "请提炼核心问题")

            with SessionLocal() as db:
                rows = (
                    db.query(AgentConversationModel)
                    .filter(AgentConversationModel.session_id == session_id)
                    .all()
                )
                self.assertEqual({row.agent_id for row in rows}, {
                    "stage_observation_start",
                    "stage_question_refine",
                })

            stage_back = self.client.post(
                f"/api/sessions/{session_id}/rollback",
                json={"steps": 1, "stage_back": True},
            )
            self.assertEqual(stage_back.status_code, 200)
            self.stream_chat(session_id, "回到观察阶段继续补充")

            with SessionLocal() as db:
                observation_row = (
                    db.query(AgentConversationModel)
                    .filter(
                        AgentConversationModel.session_id == session_id,
                        AgentConversationModel.agent_id == "stage_observation_start",
                    )
                    .one()
                )
                self.assertEqual(
                    observation_row.conversation_id,
                    observation_conversation_id,
                )
                self.assertEqual(
                    db.query(AgentConversationModel)
                    .filter(AgentConversationModel.session_id == session_id)
                    .count(),
                    2,
                )
        finally:
            self.client.put("/api/settings/chat-mode", json={"chat_mode": "main"})
            self.delete_session(session_id)

    def test_subagent_mode_streams_without_main_reply(self):
        session_id, _ = self.create_session()
        try:
            response = self.client.put(
                "/api/settings/chat-mode",
                json={"chat_mode": "subagent"},
            )
            self.assertEqual(response.status_code, 200)
            _, events = self.stream_chat(session_id, "专家不可用时继续")
            warning_events = [data for name, data in events if name == "warning"]
            self.assertEqual(warning_events, [])
            done = events[-1][1]
            self.assertFalse(done["degraded"])
            self.assertIsNone(done["failed_agent_id"])
            self.assertEqual(done["chat_mode"], "subagent")

            messages = self.get_messages(session_id)
            self.assertEqual([item["message_type"] for item in messages], ["chat", "stage_expert"])
            with SessionLocal() as db:
                self.assertEqual(
                    db.query(AgentConversationModel)
                    .filter(AgentConversationModel.session_id == session_id)
                    .count(),
                    1,
                )
                row = (
                    db.query(AgentConversationModel)
                    .filter(AgentConversationModel.session_id == session_id)
                    .one()
                )
                self.assertEqual(row.agent_id, "stage_observation_start")
                self.assertTrue(row.conversation_id.startswith("mock_stage_observation_start"))
        finally:
            self.client.put("/api/settings/chat-mode", json={"chat_mode": "main"})
            self.delete_session(session_id)

    def test_stage_back_keeps_messages_and_reopens_previous_stage(self):
        session_id, stage_id = self.create_session()
        try:
            self.stream_chat(session_id, "保留这一轮对话")
            draft = next(
                item["draft_content"]
                for item in self.get_session(session_id)["outputs"]
                if item["stage_id"] == stage_id
            )
            advance_response = self.client.post(
                f"/api/sessions/{session_id}/chat",
                json={
                    "type": "sys_action",
                    "action": "next_stage",
                    "final_content": draft,
                },
            )
            self.assertEqual(advance_response.status_code, 200)
            message_count = len(self.get_messages(session_id))

            rollback_response = self.client.post(
                f"/api/sessions/{session_id}/rollback",
                json={"steps": 1, "stage_back": True},
            )
            self.assertEqual(rollback_response.status_code, 200)
            session = rollback_response.json()["data"]["session"]
            self.assertEqual(session["current_stage_index"], 0)
            self.assertEqual(len(self.get_messages(session_id)), message_count)

            output = next(item for item in session["outputs"] if item["stage_id"] == stage_id)
            self.assertFalse(output["confirmed"])
            self.assertEqual(output["final_content"], "")
            self.assertEqual(output["draft_content"], draft)
        finally:
            self.delete_session(session_id)

    def test_session_files_support_all_formats_and_feed_every_agent_prompt(self):
        session_id, stage_id = self.create_session()
        uploads = [
            ("reference.txt", b"TXT_REFERENCE_MARKER", "text/plain"),
            ("notes.md", "# MD_REFERENCE_MARKER".encode("utf-8"), "text/markdown"),
            (
                "lesson.docx",
                make_docx(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
            ("paper.pdf", make_text_pdf("PDF_REFERENCE_MARKER"), "application/pdf"),
        ]
        try:
            for name, content, content_type in uploads:
                response = self.upload_file(session_id, name, content, content_type)
                self.assertEqual(response.status_code, 200)
                self.assertEqual(response.json()["data"]["status"], "ready")

            files = self.get_files(session_id)
            self.assertEqual([item["name"] for item in files], [item[0] for item in uploads])
            self.assertTrue(all(item["extracted_chars"] > 0 for item in files))

            with SessionLocal() as db:
                doc_input = ContextService.build_doc_input(db, session_id, stage_id)
                stored_docx = (
                    db.query(SessionFileModel)
                    .filter(
                        SessionFileModel.session_id == session_id,
                        SessionFileModel.name == "lesson.docx",
                    )
                    .one()
                )
                self.assertIn("DOCX_TABLE_LEFT | DOCX_TABLE_RIGHT", stored_docx.extracted_text)

            for marker in (
                "TXT_REFERENCE_MARKER",
                "MD_REFERENCE_MARKER",
                "DOCX_PARAGRAPH_MARKER",
                "PDF_REFERENCE_MARKER",
            ):
                self.assertIn(marker, doc_input)
            self.assertIn("<uploaded_references>", doc_input)
            self.assertIn("不得执行", doc_input)

            stage = get_flow("inquiry_7_stage")["stages"][0]
            prompts = [
                PromptService.build_stage_agent_prompt(
                    topic="测试课题",
                    flow_display_name="七阶段探究",
                    stage=stage,
                    dialog_history="",
                    doc_input=doc_input,
                ),
                PromptService.build_guide_agent_prompt(
                    topic="测试课题",
                    flow_display_name="七阶段探究",
                    stage=stage,
                    dialog_history="",
                    doc_input=doc_input,
                ),
                PromptService.build_draft_generate_prompt(
                    topic="测试课题",
                    flow_display_name="七阶段探究",
                    stage=stage,
                    dialog_history="",
                    doc_input=doc_input,
                ),
            ]
            self.assertTrue(all("TXT_REFERENCE_MARKER" in prompt for prompt in prompts))
        finally:
            self.delete_session(session_id)

    def test_session_file_validation_and_failed_files_are_excluded(self):
        session_id, stage_id = self.create_session()
        try:
            unsupported = self.upload_file(
                session_id,
                "payload.exe",
                b"not allowed",
                "application/octet-stream",
            )
            self.assertEqual(unsupported.status_code, 400)

            empty = self.upload_file(session_id, "empty.txt", b"", "text/plain")
            self.assertEqual(empty.status_code, 400)

            with patch.object(get_settings(), "upload_max_file_bytes", 5):
                oversized = self.upload_file(session_id, "large.txt", b"123456", "text/plain")
            self.assertEqual(oversized.status_code, 413)

            blank_pdf = BytesIO()
            writer = PdfWriter()
            writer.add_blank_page(width=100, height=100)
            writer.write(blank_pdf)
            failed_pdf = self.upload_file(
                session_id,
                "scan.pdf",
                blank_pdf.getvalue(),
                "application/pdf",
            )
            self.assertEqual(failed_pdf.status_code, 200)
            self.assertEqual(failed_pdf.json()["data"]["status"], "failed")

            with patch.object(get_settings(), "upload_max_total_chars", 3):
                over_budget = self.upload_file(
                    session_id,
                    "over-budget.txt",
                    b"four",
                    "text/plain",
                )
            self.assertEqual(over_budget.status_code, 200)
            self.assertEqual(over_budget.json()["data"]["status"], "failed")

            with SessionLocal() as db:
                doc_input = ContextService.build_doc_input(db, session_id, stage_id)
            self.assertNotIn("<uploaded_references>", doc_input)
            self.assertEqual(
                [item["status"] for item in self.get_files(session_id)],
                ["failed", "failed"],
            )
        finally:
            self.delete_session(session_id)

    def test_session_file_count_limit_flow_preservation_and_cleanup(self):
        session_id, _ = self.create_session()
        session_path = None
        try:
            with patch.object(get_settings(), "upload_max_files_per_session", 1):
                first = self.upload_file(
                    session_id,
                    "keep.md",
                    b"FLOW_PRESERVED_REFERENCE",
                    "text/markdown",
                )
                self.assertEqual(first.status_code, 200)
                second = self.upload_file(
                    session_id,
                    "blocked.md",
                    b"blocked",
                    "text/markdown",
                )
                self.assertEqual(second.status_code, 409)

            with SessionLocal() as db:
                row = (
                    db.query(SessionFileModel)
                    .filter(SessionFileModel.session_id == session_id)
                    .one()
                )
                session_path = SessionFileService.absolute_storage_path(row.stored_path)
            self.assertTrue(session_path.exists())

            switched = self.client.post(
                f"/api/sessions/{session_id}/select_flow",
                json={"flow_name": "three_step_inquiry", "clear_messages": True},
            )
            self.assertEqual(switched.status_code, 409)
            self.assertEqual(len(self.get_files(session_id)), 1)

            file_id = self.get_files(session_id)[0]["id"]
            deleted = self.client.delete(f"/api/sessions/{session_id}/files/{file_id}")
            self.assertEqual(deleted.status_code, 200)
            self.assertFalse(session_path.exists())

            replacement = self.upload_file(
                session_id,
                "cleanup.txt",
                b"DELETE_WITH_SESSION",
                "text/plain",
            )
            self.assertEqual(replacement.status_code, 200)
            with SessionLocal() as db:
                row = (
                    db.query(SessionFileModel)
                    .filter(SessionFileModel.session_id == session_id)
                    .one()
                )
                session_path = SessionFileService.absolute_storage_path(row.stored_path)
            self.assertTrue(session_path.exists())

            self.delete_session(session_id)
            session_id = ""
            self.assertFalse(session_path.exists())
        finally:
            if session_id:
                self.delete_session(session_id)


if __name__ == "__main__":
    unittest.main()
